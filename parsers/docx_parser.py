"""
Python-docx based parser for Word document (.docx, .doc) text extraction.

This module uses python-docx to extract text from Word documents without
requiring external APIs or heavy ML models.
"""

import os
import logging
from typing import Dict, Any

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Please install: pip install python-docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxParser:
    """
    A Word document parser using python-docx library for text extraction.

    Features:
    - Fast local DOCX/DOC parsing without external APIs
    - Extract text from paragraphs and tables
    - Simple text extraction or structured output with formatting info
    - Support for Word 2007+ documents (.docx)
    """

    def __init__(self):
        """
        Initialize the DocxParser.

        Raises:
            ImportError: If python-docx is not installed
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not available. Install it with: pip install python-docx"
            )

        self.supported_extensions = ['.docx', '.doc']

    def parse_docx(self, file_path: str, use_markdown: bool = True) -> str:
        """
        Parse DOCX file and extract text content.

        Args:
            file_path: Path to the DOCX file
            use_markdown: If True, returns simple text. If False, returns structured text with formatting

        Returns:
            Extracted text content as string

        Raises:
            FileNotFoundError: If file does not exist
            Exception: If DOCX parsing fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = Document(file_path)

            if use_markdown:
                # Simple text extraction
                text_content = ""

                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content += " | ".join(row_text) + "\n"
            else:
                # Structured extraction preserving layout
                text_content = self._extract_structured_docx(doc)

            return text_content.strip()

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise

    def parse_document(self, file_path: str, use_markdown: bool = True) -> Dict[str, Any]:
        """
        Parse a Word document and return its content with metadata.

        Args:
            file_path: Path to the DOCX file
            use_markdown: If True, returns simple text. If False, returns structured text

        Returns:
            Dictionary containing:
                - file_path: Original file path
                - file_name: File name
                - file_extension: File extension (.docx, .doc)
                - text_content: Parsed document content
                - content_length: Length of parsed content
                - word_count: Number of words in content
                - parsing_method: 'docx'
                - format_used: 'markdown' or 'structured'

        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If file is not a DOCX/DOC
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)

        if file_extension not in self.supported_extensions:
            raise ValueError(
                f"Unsupported file type: {file_extension}. "
                f"Only DOCX/DOC files are supported."
            )

        format_type = "markdown" if use_markdown else "structured"
        logger.info(f"Parsing document: {file_name} (format: {format_type})")

        # Extract text
        text_content = self.parse_docx(file_path, use_markdown)

        if not text_content:
            logger.warning(f"No text content extracted from {file_name}")

        return {
            'file_path': file_path,
            'file_name': file_name,
            'file_extension': file_extension,
            'text_content': text_content,
            'content_length': len(text_content),
            'word_count': len(text_content.split()) if text_content else 0,
            'parsing_method': 'docx',
            'format_used': format_type
        }

    def is_supported_file(self, file_path: str) -> bool:
        """Check if file format is supported (.docx or .doc)."""
        extension = os.path.splitext(file_path)[1].lower()
        return extension in self.supported_extensions

    def _extract_structured_docx(self, doc) -> str:
        """
        Extract text from DOCX preserving structure and formatting information.

        Args:
            doc: python-docx Document object

        Returns:
            Structured text with formatting markers
        """
        text_parts = []

        # Extract paragraphs with formatting info
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                # Add paragraph structure info
                text_parts.append(f"[PARAGRAPH {i+1}] {paragraph.text}")

                # Add formatting information if available
                if paragraph.style:
                    text_parts.append(f"[STYLE: {paragraph.style.name}]")

        # Extract tables with structure preservation
        for table_idx, table in enumerate(doc.tables):
            text_parts.append(f"\n=== TABLE {table_idx + 1} ===")

            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        row_data.append(f"[C{cell_idx+1}] {cell.text.strip()}")

                if row_data:
                    text_parts.append(f"[ROW {row_idx+1}] {' | '.join(row_data)}")

        return "\n".join(text_parts)


def parse_docx(
    file_path: str,
    use_markdown: bool = True,
    output_path: str = None
) -> Dict[str, Any]:
    """
    Convenience function to parse a Word document with minimal setup.

    This function provides a simple interface to parse DOCX/DOC files using python-docx.
    Fast local parsing without external APIs or ML models.

    Args:
        file_path: Path to the DOCX/DOC file
        use_markdown: Return simple text (True) or structured text with formatting (False)
        output_path: Optional path to save the parsed content

    Returns:
        Dictionary with parsed document content and metadata

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file does not exist
        ValueError: If file is not a DOCX/DOC

    Example:
        >>> from docx_parser import parse_docx
        >>> result = parse_docx("document.docx")
        >>> print(result['text_content'])
        >>>
        >>> # Save to file
        >>> result = parse_docx("document.docx", output_path="output.txt")
        >>>
        >>> # Get structured output with formatting
        >>> result = parse_docx("document.docx", use_markdown=False)
    """
    # Initialize parser
    parser = DocxParser()

    # Parse document
    result = parser.parse_document(file_path, use_markdown=use_markdown)

    # Save to file if output path provided
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result['text_content'])
        logger.info(f"Parsed content saved to: {output_path}")

    return result
